import docx
from docx import Document
from docx.shared import Cm, Mm, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx2pdf import convert



def pharmacy_reciept(pat_name, acc_addr, email_id, meds, med_price):
    doc = docx.Document()
    sections = doc.sections[0]
    sections.page_height = Cm(17)
    sections.page_width = Cm(20)
    sections.top_margin = Cm(1.27)
    sections.bottom_margin = Cm(1.27)
    sections.left_margin = Cm(1.27)
    sections.right_margin = Cm(1.27)
    P = doc.add_paragraph()
    run = P.add_run('Pharmacy Bill Receipt')
    run.font.color.rgb = RGBColor(0, 0, 51)
    run.font.size = Pt(21)
    run.bold = True
    P.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    P.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    name = 'Patient Name : '+str(pat_name)
    run1 = p.add_run(name)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    addr = 'Account Address : '+str(acc_addr)
    run1 = p.add_run(addr)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    email = 'Email Id : '+str(email_id)
    run1 = p.add_run(email)
    run1.font.size = Pt(12)

    p1 = doc.add_paragraph()
    run2 = p1.add_run('Medicine Details : ')
    run2.font.size = Pt(12)

    #meds = [[1,'Asprin',6,3],[2,'Cobiflame',5,2]]
    table = doc.add_table(rows=1, cols=4)
    row = table.rows[0].cells
    row[0].text = 'S.No'
    row[1].text = 'Med Name'
    row[2].text = 'Qty'
    row[3].text = 'Price(ETH)'
    for sno, name, qty, price in meds:
        row = table.add_row().cells
        row[0].text = str(sno)
        row[1].text = name
        row[2].text = str(qty)
        row[3].text = str(price)
    #table.style = 'Colorful List'
    table.style = 'Table Grid'

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(12)
    totprice = 'Total Amount : '+str(med_price)+' ETH '
    run2 = p3.add_run(totprice)
    run2.font.size = Pt(12)
    run2.add_break()
    run2.add_break()

    amtpaid = 'Pharmacy Bill Amount : '+'Paid '
    run2 = p3.add_run(amtpaid)
    run2 = p3.add_run('✔')
    run2.font.size = Pt(12)
    run2.add_break()
    run2.add_break()

    #doc.add_heading('Image with Defined Size:', 3)
    p4 = doc.add_paragraph()
    run3 = p4.add_run('Pharmacist Sign')
    run3.font.color.rgb = RGBColor(97, 6, 6)
    run3.font.size = Pt(10)
    run3.bold = True
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_picture('phar-sign.png', width=Cm(3), height=Cm(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save('phar.docx')
    convert('phar.docx')
    


def basicpay_reciept(pat_name, acc_addr, email_id, insr_plan, insr_availdate, insr_pay):
    doc = docx.Document()
    sections = doc.sections[0]
    sections.page_height = Cm(17)
    sections.page_width = Cm(20)
    sections.top_margin = Cm(1.27)
    sections.bottom_margin = Cm(1.27)
    sections.left_margin = Cm(1.27)
    sections.right_margin = Cm(1.27)
    P = doc.add_paragraph()
    run = P.add_run('Insurance Plan Avail Reciept')
    run.font.color.rgb = RGBColor(0, 0, 51)
    run.font.size = Pt(21)
    run.bold = True
    P.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    P.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    name = 'Patient Name : '+str(pat_name)
    run1 = p.add_run(name)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    addr = 'Account Address : '+str(acc_addr)
    run1 = p.add_run(addr)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    email = 'Email Id : '+str(email_id)
    run1 = p.add_run(email)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    plan = 'Insurance Plan : '+str(insr_plan)
    run1 = p.add_run(plan)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    insrdate = 'Insurance Avail Date : '+str(insr_availdate)
    run1 = p.add_run(insrdate)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    amtpaid = 'Insurance Plan Avail Amount : '+str(insr_pay)+' ETH '
    run1 = p.add_run(amtpaid)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    #doc.add_heading('Image with Defined Size:', 3)
    p1 = doc.add_paragraph()
    run2 = p1.add_run('Insurance Company Sign')
    run2.font.color.rgb = RGBColor(97, 6, 6)
    run2.font.size = Pt(10)
    run2.bold = True
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_picture('ins-sign.png', width=Cm(3), height=Cm(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save('insrbasicpay.docx')
    convert('insrbasicpay.docx')
    


def premium_reciept(pat_name, acc_addr, email_id, insr_plan, premamt, insrval):
    doc = docx.Document()
    sections = doc.sections[0]
    sections.page_height = Cm(17)
    sections.page_width = Cm(20)
    sections.top_margin = Cm(1.27)
    sections.bottom_margin = Cm(1.27)
    sections.left_margin = Cm(1.27)
    sections.right_margin = Cm(1.27)
    P = doc.add_paragraph()
    run = P.add_run('Insurance Premium Reciept')
    run.font.color.rgb = RGBColor(0, 0, 51)
    run.font.size = Pt(21)
    run.bold = True
    P.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    P.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    name = 'Patient Name : '+str(pat_name)
    run1 = p.add_run(name)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    addr = 'Account Address : '+str(acc_addr)
    run1 = p.add_run(addr)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    email = 'Email Id : '+str(email_id)
    run1 = p.add_run(email)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    plan = 'Insurance Plan : '+str(insr_plan)
    run1 = p.add_run(plan)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    premium = 'Premium : '+str(premamt)
    run1 = p.add_run(premium)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    premium_paid = 'Premium Bill Amount : '+'Paid '
    run1 = p.add_run(premium_paid)
    run1 = p.add_run('✔')
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    insr_val = 'Insurance Plan Validity : '+str(insrval)
    run1 = p.add_run(insr_val)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    #doc.add_heading('Image with Defined Size:', 3)
    p1 = doc.add_paragraph()
    run2 = p1.add_run('Insurance Company Sign')
    run2.font.color.rgb = RGBColor(97, 6, 6)
    run2.font.size = Pt(10)
    run2.bold = True
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_picture('ins-sign.png', width=Cm(3), height=Cm(2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save('premiumpay.docx')
    convert('premiumpay.docx')


def fullamount_reciept(from_pat, pat_acc, acc, insr_plan, totcost, remamt):
    doc = docx.Document()
    sections = doc.sections[0]
    sections.page_height = Cm(17)
    sections.page_width = Cm(22)
    sections.top_margin = Cm(1.27)
    sections.bottom_margin = Cm(1.27)
    sections.left_margin = Cm(1.27)
    sections.right_margin = Cm(1.27)
    P = doc.add_paragraph()
    run = P.add_run('Insurance Full Amount Reciept')
    run.font.color.rgb = RGBColor(0, 0, 51)
    run.font.size = Pt(21)
    run.bold = True
    P.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    P.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    insracc = 'Insurance Company Account : '+str(acc)
    run1 = p.add_run(insracc)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    name = 'Patient Name : '+str(from_pat)
    run1 = p.add_run(name)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    patacc = 'Patient Account: '+str(pat_acc)
    run1 = p.add_run(patacc)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    plan = 'Insurance Plan : '+str(insr_plan)
    run1 = p.add_run(plan)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    rem_amt = 'Remaining Amount : '+str(remamt)+' ETH '
    run1 = p.add_run(rem_amt)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    rem_amt_paid = 'Remaining Bill Amount : '+'Paid '
    run1 = p.add_run(rem_amt_paid)
    run1 = p.add_run('✔')
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    totalcost = 'Total Cost : '+str(totcost)+' ETH '
    run1 = p.add_run(totalcost)
    run1.font.size = Pt(12)
    run1.add_break()
    run1.add_break()

    #doc.add_heading('Image with Defined Size:', 3)
    p1 = doc.add_paragraph()
    run2 = p1.add_run('Hospital Admin Sign\t\t\t\t')
    run2.font.color.rgb = RGBColor(97, 6, 6)
    run2.font.size = Pt(11)
    run2.bold = True
    #p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run2 = p1.add_run('Lab Admin Sign\t\t\t\t')
    run2.font.color.rgb = RGBColor(97, 6, 6)
    run2.font.size = Pt(11)
    run2.bold = True
    #p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run2 = p1.add_run('Pharmacist Sign\t\t\t\t')
    run2.font.color.rgb = RGBColor(97, 6, 6)
    run2.font.size = Pt(11)
    run2.bold = True
    #p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p2 = doc.add_paragraph()
    #p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run3 = p2.add_run()
    run3.add_picture('hos-sign.png', width=Cm(3.5), height=Cm(2))
    run3 = p2.add_run('\t\t\t\t')
    run3.add_picture('lab-sign.png', width=Cm(3.5), height=Cm(2))
    run3 = p2.add_run('\t\t\t\t')
    run3.add_picture('phar-sign.png', width=Cm(3.5), height=Cm(2))
    '''
    last_paragraph = doc.paragraphs[-2] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    '''
    doc.save('fullpay.docx')
    convert('fullpay.docx')

# fullamount_reciept('Ben','0x90F4B9227EdBd11cC97C0DEB5D8Bc98385119Cf7','0x80C9715c92Fa6B377275777e7870627a1F1b0aA8','A','22','10')
